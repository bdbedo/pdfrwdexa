#!/usr/bin/env python3
import sys
from io import BytesIO
from pathlib import Path
from typing import List, Optional

import fitz
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _safe_text(block: dict) -> str:
    return ' '.join(span.get('text', '').strip() for line in block.get('lines', []) for span in line.get('spans', []) if span.get('text', '').strip()).strip()


def _extract_lines(block: dict) -> List[dict]:
    lines = []
    for line in block.get('lines', []):
        spans = []
        for span in line.get('spans', []):
            text = span.get('text', '').strip()
            if text:
                spans.append({
                    'text': text,
                    'font': span.get('font', ''),
                    'size': float(span.get('size', 11)),
                    'flags': int(span.get('flags', 0)),
                })
        if spans:
            lines.append({
                'spans': spans,
                'bbox': line.get('bbox', [0, 0, 0, 0]),
            })
    return lines


def _is_probable_heading(text: str, font_size: float, block_bbox: list, page_height: float, previous_bbox: Optional[list], previous_size: Optional[float]) -> bool:
    if not text or len(text.split()) > 12:
        return False
    if font_size >= 14:
        return True
    if block_bbox[1] < page_height * 0.2 and font_size >= 12.5:
        return True
    if previous_bbox and previous_size and previous_size >= 13 and font_size >= 12:
        return True
    return False


def _detect_list_type(text: str) -> Optional[str]:
    if text.startswith(('-', '•')):
        return 'bullet'
    if text[:2].isdigit() and text[1:2] in {'.', ')'}:
        return 'numbered'
    return None


def _build_table(doc: Document, table_data: List[List[str]]) -> None:
    if not table_data:
        return
    table = doc.add_table(rows=1, cols=max(len(row) for row in table_data))
    table.style = 'Table Grid'
    for row_index, row in enumerate(table_data):
        cells = table.rows[row_index].cells
        for col_index, content in enumerate(row):
            if col_index >= len(cells):
                break
            cells[col_index].text = content


def build_docx_from_pdf(pdf_path: Path, output_path: Path) -> None:
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    try:
        pdf_document = fitz.open(pdf_path)
    except Exception as exc:
        raise RuntimeError(f'Unable to read PDF: {exc}') from exc

    if pdf_document.is_encrypted:
        raise RuntimeError('This PDF is password-protected. Please unlock the PDF before converting it.')

    page_count = len(pdf_document)
    if page_count == 0:
        raise RuntimeError('The PDF appears to be empty.')

    has_selectable_text = False
    for page_index in range(page_count):
        page = pdf_document[page_index]
        text_blocks = page.get_text('dict').get('blocks', [])
        for block in text_blocks:
            if block.get('type', 0) == 0 and _safe_text(block):
                has_selectable_text = True
                break
        if has_selectable_text:
            break

    if not has_selectable_text:
        raise RuntimeError('This PDF appears to be scanned or image-only. OCR support is not available in this environment yet. Please use a selectable-text PDF or add OCR support.')

    for page_index in range(page_count):
        page = pdf_document[page_index]
        page_dict = page.get_text('dict')
        blocks = page_dict.get('blocks', [])

        if page_index > 0:
            doc.add_page_break()

        previous_block_bbox = None
        previous_font_size = None
        for block in blocks:
            if block.get('type', 0) != 0:
                continue

            text = _safe_text(block)
            if not text:
                continue

            lines = _extract_lines(block)
            if not lines:
                continue

            font_sizes = [span['size'] for line in lines for span in line['spans']]
            average_size = sum(font_sizes) / len(font_sizes) if font_sizes else 11
            block_bbox = block.get('bbox', [0, 0, 0, 0])
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)

            block_y = block_bbox[1]
            if previous_block_bbox and (previous_block_bbox[1] - block_bbox[1]) > 40:
                paragraph.paragraph_format.space_before = Pt(10)
            else:
                paragraph.paragraph_format.space_before = Pt(0)

            if _detect_list_type(text):
                paragraph.style = 'List Bullet' if _detect_list_type(text) == 'bullet' else 'List Number'
                paragraph.paragraph_format.left_indent = Pt(18)
                paragraph.add_run(text.lstrip('-•').lstrip())
            else:
                run = paragraph.add_run(text)
                run.font.size = Pt(max(10.5, round(average_size * 0.9, 1)))
                if _is_probable_heading(text, average_size, block_bbox, page.rect.height, previous_block_bbox, previous_font_size):
                    run.bold = True
                    paragraph.style = 'Heading 1' if average_size >= 15 else 'Heading 2'
                    paragraph.paragraph_format.space_before = Pt(10)
                    paragraph.paragraph_format.space_after = Pt(4)

            previous_block_bbox = block_bbox
            previous_font_size = average_size

        if page.get_images():
            for img in page.get_images(full=True):
                try:
                    xref = img[0]
                    image = pdf_document.extract_image(xref)
                    image_bytes = image['image']
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(image_bytes), width=Inches(2.5))
                except Exception:
                    continue

        try:
            table_finder = page.find_tables()
            tables = list(table_finder) if hasattr(table_finder, '__iter__') else []
            for table in tables:
                rows = []
                for row in table.extract():
                    rows.append([str(cell).strip() for cell in row])
                if rows:
                    _build_table(doc, rows)
        except Exception:
            pass

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit('Usage: convert_pdf_to_docx.py <input.pdf> <output.docx>')

    input_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    build_docx_from_pdf(input_path, output_path)


if __name__ == '__main__':
    main()
